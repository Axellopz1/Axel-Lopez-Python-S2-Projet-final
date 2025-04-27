# --- Imports nécessaires ---
from docx import Document  # Pour créer et manipuler des documents Word
from docx.shared import Pt, Inches  # Pour gérer la taille des polices et des images
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # Pour aligner du texte (ex: centré)
import matplotlib.pyplot as plt  # Pour créer des graphiques
import pandas as pd  # Pour manipuler des DataFrames
import io  # Pour manipuler des flux d'images en mémoire

# --- Fonction principale pour générer la fiche Word ---
def generer_fiche_portefeuille(portefeuille, projections, stress_results, preferences, montant_investi):
    doc = Document()  # Crée un nouveau document Word vide

    # --- Titre principal du document ---
    titre = doc.add_heading('Fiche de Portefeuille Client', level=1)  # Titre de niveau 1
    titre.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centrer le titre

    doc.add_paragraph()  # Ajouter un espace

    # --- Section Profil Investisseur ---
    doc.add_heading('Profil de l\'investisseur', level=2)  # Titre de niveau 2
    doc.add_paragraph(f"Objectif Financier : {preferences.objectif}")  # Ajouter l'objectif
    doc.add_paragraph(f"Horizon d'Investissement : {preferences.horizon} ans")  # Ajouter l'horizon
    doc.add_paragraph(f"Tolérance au Risque : {preferences.tolerance}")  # Ajouter la tolérance au risque
    doc.add_paragraph(f"Montant investi : {montant_investi:,} $")  # Ajouter le montant investi

    doc.add_paragraph()  # Ajouter un espace

    # --- Petit résumé commercial automatique ---
    doc.add_heading('Résumé général', level=2)  # Titre de niveau 2
    paragraphe = (
        f"Votre portefeuille vise une croissance dynamique tout en maîtrisant les risques. "
        f"Il est diversifié entre plusieurs types d'actifs, couvrant différentes zones géographiques et secteurs économiques. "
        f"Cette allocation permet d'équilibrer rendement potentiel et gestion prudente des risques."
    )
    doc.add_paragraph(paragraphe)  # Ajouter le texte de résumé

    doc.add_paragraph()  # Ajouter un espace

    # --- Section Composition par Type d'Actif ---
    doc.add_heading('Composition par Type d\'Actif', level=2)  # Titre de niveau 2
    df = portefeuille.composition()  # Récupérer la composition du portefeuille
    if not df.empty:  # Si le portefeuille n'est pas vide
        composition = df.groupby("Type")["Valeur Totale"].sum()  # Regrouper par Type
        table = doc.add_table(rows=1, cols=2)  # Créer une table Word
        table.style = 'Light List Accent 1'  # Style visuel

        # Remplir l'entête
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Type d\'Actif'
        hdr_cells[1].text = 'Valeur Totale ($)'

        # Remplir les données de composition
        for idx, val in composition.items():
            row_cells = table.add_row().cells
            row_cells[0].text = str(idx)
            row_cells[1].text = f"{val:,.2f}"

    doc.add_paragraph()  # Ajouter un espace

    # --- Section Répartition Sectorielle ---
    doc.add_heading('Répartition Sectorielle', level=2)  # Titre
    secteur = df.groupby("Secteur")["Valeur Totale"].sum()  # Regrouper par Secteur
    table = doc.add_table(rows=1, cols=2)  # Nouvelle table
    table.style = 'Light List Accent 2'  # Style différent

    # Remplir l'entête
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Secteur'
    hdr_cells[1].text = 'Valeur Totale ($)'

    # Remplir les lignes du tableau
    for idx, val in secteur.items():
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = f"{val:,.2f}"

    doc.add_paragraph()  # Ajouter un espace

    # --- Section Répartition Géographique ---
    doc.add_heading('Répartition Géographique', level=2)  # Titre
    geo = df.groupby("Zone Géographique")["Valeur Totale"].sum()  # Regrouper par zone
    table = doc.add_table(rows=1, cols=2)  # Nouvelle table
    table.style = 'Light List Accent 3'  # Autre style

    # Remplir l'entête
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Zone Géographique'
    hdr_cells[1].text = 'Valeur Totale ($)'

    # Remplir les données
    for idx, val in geo.items():
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = f"{val:,.2f}"

    doc.add_paragraph()  # Ajouter un espace

    # --- Section Analyse Quantitative ---
    doc.add_heading('Analyse Quantitative', level=2)  # Titre
    valeur_initiale = projections.iloc[0]  # Valeur de départ
    valeur_finale = projections.iloc[-1]  # Valeur finale

    # Calculs quantitatifs
    rendement_total = ((valeur_finale - valeur_initiale) / valeur_initiale) * 100
    cagr = ((valeur_finale / valeur_initiale) ** (1 / 2) - 1) * 100  # CAGR sur 2 ans
    volatilite_mensuelle = projections.pct_change().std() * 100  # Volatilité
    max_drawdown = (projections.cummax() - projections).max() / projections.cummax().max() * 100  # Drawdown max

    # Regrouper dans un dictionnaire
    analyse_data = {
        "Rendement Total (%)": rendement_total,
        "CAGR (%)": cagr,
        "Volatilité Mensuelle (%)": volatilite_mensuelle.mean(),
        "Maximum Drawdown (%)": max_drawdown.mean(),
    }

    # Créer la table
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light List Accent 4'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Indicateur'
    hdr_cells[1].text = 'Valeur'

    # Remplir la table
    for idx, val in analyse_data.items():
        if isinstance(val, pd.Series):
            val = val.mean()  # Si jamais c'est une série
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = f"{val:.2f}"

    doc.add_paragraph()  # Ajouter un espace

    # --- Section Stress Tests ---
    doc.add_heading('Stress Tests - Pertes Maximales', level=2)
    pertes = {}

    # Calculer les pertes pour chaque scénario
    for col in stress_results.columns:
        valeur_init = stress_results[col].iloc[0]
        valeur_fin = stress_results[col].iloc[-1]
        perte = ((valeur_init - valeur_fin) / valeur_init) * 100
        pertes[col] = perte

    # Créer la table
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light List Accent 5'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Scénario'
    hdr_cells[1].text = 'Perte (%)'

    # Remplir les données
    for idx, val in pertes.items():
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = f"{val:.2f}"

    doc.add_paragraph()  # Ajouter un espace

    # --- Graphique de Répartition par Type d'Actif ---
    doc.add_heading('Graphique de Répartition', level=2)  # Titre
    fig, ax = plt.subplots()
    df.groupby("Type")["Valeur Totale"].sum().plot.pie(autopct='%1.1f%%', ax=ax)  # Pie chart
    ax.set_ylabel("")  # Pas de label sur l'axe Y
    ax.set_title("Répartition par Type d'Actif")  # Titre du graphique

    img_stream = io.BytesIO()  # Stocker l'image en mémoire
    fig.savefig(img_stream, format='png')
    img_stream.seek(0)

    from docx.shared import Inches  # (Déjà importé en haut, mais redemandé ici)
    doc.add_picture(img_stream, width=Inches(5))  # Ajouter l'image au Word

    # --- Sauvegarder le document Word ---
    doc.save("fiche_portefeuille.docx")
